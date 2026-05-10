"""PARITY-01 DOCX extractor tests.

Builds the fixture programmatically via python-docx itself (no
checked-in binary). Embeds one inline image at >=64x64 so the
icon-noise floor in _common.py admits it.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest


def _make_small_png(rgb: tuple[int, int, int], size: int = 80) -> bytes:
    from PIL import Image

    img = Image.new("RGB", (size, size), color=rgb)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture
def sample_docx_with_image(tmp_path: Path) -> tuple[Path, int]:
    """Build a DOCX with text and 1 embedded image. Returns (path, image_count)."""
    from docx import Document

    doc = Document()
    doc.add_heading("Nyra Sample DOCX", level=1)
    doc.add_paragraph(
        "This is a paragraph that should be extracted by python-docx. "
        "It mentions a specific phrase: pineapple-bridge."
    )
    doc.add_paragraph("A second paragraph for line-count assertions.")

    # Embed one inline picture.
    png_bytes = _make_small_png((0, 0, 255), size=80)
    doc.add_picture(io.BytesIO(png_bytes))

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out, 1


def test_extract_docx_text_and_image(
    tmp_project_dir: Path, sample_docx_with_image
) -> None:
    from nyrahost.extractors.docx import extract_docx

    docx_path, expected_image_count = sample_docx_with_image
    text, refs = extract_docx(
        docx_path, project_saved=tmp_project_dir / "Saved"
    )
    assert isinstance(text, str)
    assert "pineapple-bridge" in text
    assert "Nyra Sample DOCX" in text
    assert isinstance(refs, list)
    assert len(refs) == expected_image_count
    for ref in refs:
        assert ref.kind == "image"
        assert Path(ref.path).exists()


def test_extract_docx_malformed_raises(
    tmp_path: Path, tmp_project_dir: Path
) -> None:
    from nyrahost.extractors.docx import extract_docx

    bad = tmp_path / "not-a-docx.docx"
    bad.write_bytes(b"this is plain text, not a zip")
    with pytest.raises(ValueError):
        extract_docx(bad, project_saved=tmp_project_dir / "Saved")


def test_extract_docx_empty_zip_raises(
    tmp_path: Path, tmp_project_dir: Path
) -> None:
    """An empty .docx (zero bytes) must raise ValueError, not crash."""
    from nyrahost.extractors.docx import extract_docx

    bad = tmp_path / "empty.docx"
    bad.write_bytes(b"")
    with pytest.raises(ValueError):
        extract_docx(bad, project_saved=tmp_project_dir / "Saved")


def test_extract_docx_text_grows_with_input(tmp_path: Path, tmp_project_dir: Path) -> None:
    """Sanity: more paragraphs in the source -> longer extracted text."""
    from docx import Document

    from nyrahost.extractors.docx import extract_docx

    short = tmp_path / "short.docx"
    long_ = tmp_path / "long.docx"

    short_doc = Document()
    short_doc.add_paragraph("one short line.")
    short_doc.save(str(short))

    long_doc = Document()
    for i in range(20):
        long_doc.add_paragraph(
            f"Line number {i}: this is a longer paragraph for the size test."
        )
    long_doc.save(str(long_))

    short_text, _ = extract_docx(short, project_saved=tmp_project_dir / "Saved")
    long_text, _ = extract_docx(long_, project_saved=tmp_project_dir / "Saved")
    assert len(long_text) > len(short_text)
