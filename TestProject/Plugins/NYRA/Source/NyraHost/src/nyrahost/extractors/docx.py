"""DOCX extractor — python-docx 1.2.0 (LOCKED-06 approved).

Returns ``(text, list[AttachmentRef])``. Text comes from iterating
``doc.paragraphs``; embedded images come from
``doc.part.related_parts`` filtered by content_type starting with
``image/``. Pre-checks the OOXML zip's total uncompressed size against
the 100 MB cap (zip-bomb mitigation).
"""
from __future__ import annotations

from pathlib import Path
from typing import Tuple

from nyrahost.attachments import AttachmentRef
from nyrahost.extractors._common import (
    assert_ooxml_within_zip_budget,
    ingest_image_bytes,
)


def extract_docx(
    src: Path, *, project_saved: Path
) -> Tuple[str, list[AttachmentRef]]:
    assert_ooxml_within_zip_budget(src)

    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as e:
        raise ValueError(
            "python-docx not installed — cannot extract DOCX documents"
        ) from e

    try:
        doc = Document(str(src))
    except PackageNotFoundError as e:
        raise ValueError(f"malformed DOCX {src.name}: {e}") from e
    except Exception as e:  # python-docx wraps lxml errors loosely
        raise ValueError(f"unreadable DOCX {src.name}: {e}") from e

    text_chunks: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            text_chunks.append(para.text)
    # Tables are common in design docs and python-docx exposes them
    # separately from paragraphs — skipping them would silently drop
    # half the content of a typical "spec sheet" DOCX.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_chunks.append(cell.text)

    image_refs: list[AttachmentRef] = []
    for rel in doc.part.related_parts.values():
        content_type = getattr(rel, "content_type", "") or ""
        if not content_type.startswith("image/"):
            continue
        blob = getattr(rel, "blob", None)
        if not blob:
            continue
        ref = ingest_image_bytes(blob, project_saved=project_saved)
        if ref is not None:
            image_refs.append(ref)

    text = "\n".join(text_chunks)
    return text, image_refs
